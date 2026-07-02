from flask import Flask, render_template, request, send_file
from PyPDF2 import PdfMerger
from pypdf import PdfReader, PdfWriter
from docx import Document
from PIL import Image
import pdfplumber
import pikepdf
import pandas as pd
import cv2
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors

app = Flask(__name__)

# ---------------- HOME ----------------
@app.route("/")
def home():
    return render_template("index.html")
    @app.route("/about")
def about():
    return render_template("about.html")

@app.route("/contact")
def contact():
    return render_template("contact.html")

@app.route("/privacy")
def privacy():
    return render_template("privacy.html")


@app.route("/pdf")
def pdf():
    return render_template("pdf.html")


@app.route("/excel")
def excel():
    return render_template("excel.html")


@app.route("/word")
def word():
    return render_template("word.html")


@app.route("/scanner")
def scanner():
    return render_template("scanner.html")


# ---------------- PDF TOOLS ----------------

@app.route("/merge-pdf", methods=["POST"])
def merge_pdf():
    files = request.files.getlist("pdfs")
    merger = PdfMerger()

    for file in files:
        merger.append(file)

    output = "merged.pdf"
    merger.write(output)
    merger.close()

    return send_file(output, as_attachment=True)


@app.route("/split-pdf", methods=["POST"])
def split_pdf():
    file = request.files["pdf"]
    page = int(request.form["split_page"])

    reader = PdfReader(file)
    writer = PdfWriter()

    for i in range(page):
        writer.add_page(reader.pages[i])

    output = "split.pdf"

    with open(output, "wb") as f:
        writer.write(f)

    return send_file(output, as_attachment=True)


@app.route("/compress-pdf", methods=["POST"])
def compress_pdf():
    file = request.files["pdf"]

    input_path = "input.pdf"
    output_path = "compressed.pdf"

    file.save(input_path)

    pdf = pikepdf.open(input_path)
    pdf.save(output_path, compress_streams=True)
    pdf.close()

    return send_file(output_path, as_attachment=True)


# ---------------- WORD TOOLS ----------------

@app.route("/pdf-to-word", methods=["POST"])
def pdf_to_word():
    file = request.files["pdf"]

    pdf_path = "file.pdf"
    docx_path = "file.docx"

    file.save(pdf_path)

    doc = Document()

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)

    doc.save(docx_path)

    return send_file(docx_path, as_attachment=True)


@app.route("/word-to-pdf", methods=["POST"])
def word_to_pdf():
    file = request.files["docx"]

    docx_path = "input.docx"
    pdf_path = "output.pdf"

    file.save(docx_path)

    from docx2pdf import convert
    convert(docx_path, pdf_path)

    return send_file(pdf_path, as_attachment=True)


@app.route("/docx-to-txt", methods=["POST"])
def docx_to_txt():
    file = request.files["docx"]

    docx_path = "input.docx"
    txt_path = "output.txt"

    file.save(docx_path)

    doc = Document(docx_path)
    text = "\n".join([p.text for p in doc.paragraphs])

    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(text)

    return send_file(txt_path, as_attachment=True)


# ---------------- EXCEL TOOLS ----------------

@app.route("/csv-to-excel", methods=["POST"])
def csv_to_excel():
    file = request.files["csv"]

    df = pd.read_csv(file)
    output = "output.xlsx"

    df.to_excel(output, index=False)

    return send_file(output, as_attachment=True)


@app.route("/excel-to-pdf", methods=["POST"])
def excel_to_pdf():
    file = request.files["excel"]

    df = pd.read_excel(file)

    output = "excel.pdf"
    doc = SimpleDocTemplate(output)

    data = [df.columns.tolist()] + df.values.tolist()

    table = Table(data)

    style = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
    ])

    table.setStyle(style)
    doc.build([table])

    return send_file(output, as_attachment=True)


@app.route("/remove-duplicates", methods=["POST"])
def remove_duplicates():
    file = request.files["excel"]

    df = pd.read_excel(file)
    df = df.drop_duplicates()

    output = "cleaned.xlsx"
    df.to_excel(output, index=False)

    return send_file(output, as_attachment=True)


# ---------------- IMAGE TOOLS ----------------

@app.route("/image-to-pdf", methods=["POST"])
def image_to_pdf():
    file = request.files["image"]

    img = Image.open(file).convert("RGB")
    output = "image.pdf"

    img.save(output)

    return send_file(output, as_attachment=True)


# ---------------- SCANNER ----------------

@app.route("/scan-document", methods=["POST"])
def scan_document():
    file = request.files["image"]

    input_img = "scan.jpg"
    output_img = "scanned.jpg"
    output_pdf = "scanned.pdf"

    file.save(input_img)

    image = cv2.imread(input_img)

    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    gray = cv2.GaussianBlur(gray, (5, 5), 0)

    edges = cv2.Canny(gray, 75, 200)

    contours, _ = cv2.findContours(edges, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)
    contours = sorted(contours, key=cv2.contourArea, reverse=True)

    cropped = image

    for c in contours:
        approx = cv2.approxPolyDP(c, 0.02 * cv2.arcLength(c, True), True)

        if len(approx) == 4:
            x, y, w, h = cv2.boundingRect(approx)
            cropped = image[y:y+h, x:x+w]
            break

    cv2.imwrite(output_img, cropped)

    img = Image.open(output_img).convert("RGB")
    img.save(output_pdf)

    return send_file(output_pdf, as_attachment=True)


# ---------------- RUN APP ----------------

if __name__ == "__main__":
    app.run(debug=True)
